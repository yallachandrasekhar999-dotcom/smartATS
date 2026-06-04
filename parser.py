import os
import re
import spacy
import nltk
import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

# NLP Setup & Asset checks
def setup_nlp():
    print("Checking and downloading required NLP models...")
    # NLTK resources
    nltk_resources = {
        'stopwords': 'corpora/stopwords',
        'punkt': 'tokenizers/punkt',
        'wordnet': 'corpora/wordnet',
        'omw-1.4': 'corpora/omw-1.4'
    }
    for res_name, res_path in nltk_resources.items():
        try:
            nltk.data.find(res_path)
            print(f"NLTK resource '{res_name}' is already installed.")
        except LookupError:
            print(f"Downloading NLTK resource '{res_name}'...")
            nltk.download(res_name, quiet=True)

    # spaCy model check
    try:
        spacy.load('en_core_web_sm')
        print("spaCy 'en_core_web_sm' model is already installed.")
    except OSError:
        print("Downloading spaCy 'en_core_web_sm' model...")
        import subprocess
        import sys
        subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)

# Run NLP setup on initialization
setup_nlp()

# Load spaCy NLP model
nlp = spacy.load('en_core_web_sm')
from nltk.corpus import stopwords
STOPWORDS = set(stopwords.words('english'))

# Large list of common industry skills grouped by category
SKILL_DICTIONARY = {
    'Languages': [
        'python', 'java', 'c++', 'c#', 'javascript', 'typescript', 'ruby', 'golang', 'go',
        'rust', 'php', 'swift', 'kotlin', 'html', 'html5', 'css', 'css3', 'sql', 'r', 'bash', 
        'scala', 'perl', 'dart', 'matlab'
    ],
    'Frameworks & Libraries': [
        'react', 'react.js', 'reactjs', 'angular', 'angularjs', 'vue', 'vue.js', 'vuejs', 
        'node', 'node.js', 'nodejs', 'express', 'express.js', 'django', 'flask', 'fastapi', 
        'spring', 'spring boot', 'asp.net', 'laravel', 'rails', 'ruby on rails', 'jquery', 
        'bootstrap', 'tailwind', 'tailwindcss', 'next.js', 'nextjs', 'nuxt.js', 'svelte',
        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'opencv',
        'redux', 'graphql', 'apollo', 'hibernate', 'sequelize'
    ],
    'Databases': [
        'mysql', 'postgresql', 'postgres', 'mongodb', 'sqlite', 'redis', 'oracle', 
        'dynamodb', 'cassandra', 'mariadb', 'firebase', 'elasticsearch', 'neo4j', 'mssql'
    ],
    'Cloud & DevOps': [
        'aws', 'amazon web services', 'azure', 'gcp', 'google cloud', 'google cloud platform', 
        'docker', 'kubernetes', 'k8s', 'ci/cd', 'jenkins', 'github actions', 'gitlab ci',
        'git', 'github', 'gitlab', 'terraform', 'ansible', 'linux', 'unix', 'heroku', 
        'nginx', 'apache', 'aws ec2', 'aws s3', 'rds', 'lambda', 'circleci'
    ],
    'Methodologies & Tools': [
        'agile', 'scrum', 'jira', 'confluence', 'gitflow', 'rest api', 'soap', 'graphql',
        'microservices', 'mvc', 'oop', 'tdd', 'unit testing', 'figma', 'postman', 
        'powerbi', 'tableau', 'excel', 'word', 'powerpoint'
    ],
    'Soft Skills & Others': [
        'communication', 'leadership', 'teamwork', 'problem solving', 'critical thinking',
        'time management', 'adaptability', 'creativity', 'project management', 'collaboration',
        'interpersonal', 'negotiation', 'analytical', 'customer service'
    ]
}

# Flatten list for regex search
ALL_FLAT_SKILLS = []
for cat, skills in SKILL_DICTIONARY.items():
    ALL_FLAT_SKILLS.extend(skills)
ALL_FLAT_SKILLS = sorted(list(set(ALL_FLAT_SKILLS)), key=len, reverse=True)


def extract_text_from_pdf(pdf_path):
    """Extracts raw text from a PDF file using pdfplumber, falling back to PyPDF2."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}. Trying PyPDF2...")
        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as ex:
            print(f"PyPDF2 failed: {ex}")
    return text


def extract_text_from_docx(docx_path):
    """Extracts raw text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"python-docx failed: {e}")
    return text


def extract_text(file_path):
    """Extracts text depending on the file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        return ""


def extract_contact_info(text):
    """Uses Regex and NLP to extract Email and Phone from text."""
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    email_match = re.search(email_pattern, text)
    phone_match = re.search(phone_pattern, text)
    
    email = email_match.group(0) if email_match else None
    phone = phone_match.group(0) if phone_match else None
    
    return email, phone


def extract_name(text):
    """Attempts to extract a name using spaCy NER or first line heuristics."""
    doc = nlp(text[:800]) # Scan the first 800 characters for name
    
    # Try PERSON entity
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            # Basic validation: Name should be 2-3 words
            name = ent.text.strip().replace('\n', ' ')
            if len(name.split()) >= 2 and len(name.split()) <= 4:
                return name
                
    # Heuristic fallback: Use the first non-empty line if it looks like a name
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:3]:
        # Exclude contact info or metadata lines
        if '@' not in line and not re.search(r'\d{4}', line) and len(line.split()) >= 2 and len(line.split()) <= 4:
            return line
            
    return "Candidate Name Not Found"



def format_skill_name(skill_name):
    skill_lower = skill_name.lower().strip()
    
    # Specific casing mapping
    custom_map = {
        'postgresql': 'PostgreSQL',
        'mysql': 'MySQL',
        'mongodb': 'MongoDB',
        'sqlite': 'SQLite',
        'neo4j': 'Neo4j',
        'mssql': 'MSSQL',
        'dynamodb': 'DynamoDB',
        'mariadb': 'MariaDB',
        'github': 'GitHub',
        'gitlab': 'GitLab',
        'graphql': 'GraphQL',
        'ci/cd': 'CI/CD',
        'c++': 'C++',
        'c#': 'C#',
        'aws': 'AWS',
        'gcp': 'GCP',
        'mvc': 'MVC',
        'oop': 'OOP',
        'tdd': 'TDD',
        'api': 'API',
        'rest api': 'REST API',
        'soap': 'SOAP',
        'html': 'HTML',
        'html5': 'HTML5',
        'css': 'CSS',
        'css3': 'CSS3',
        'sql': 'SQL',
        'r': 'R',
        'powerbi': 'PowerBI',
        'docx': 'DOCX',
        'pdf': 'PDF'
    }
    
    if skill_lower in custom_map:
        return custom_map[skill_lower]
        
    if skill_lower in ['react', 'node', 'vue', 'next', 'nuxt', 'express']:
        return skill_lower.title() + ".js"
        
    if skill_lower in ['reactjs', 'react.js']:
        return 'React.js'
    if skill_lower in ['nodejs', 'node.js']:
        return 'Node.js'
    if skill_lower in ['vuejs', 'vue.js']:
        return 'Vue.js'
    if skill_lower in ['nextjs', 'next.js']:
        return 'Next.js'
    if skill_lower in ['nuxtjs', 'nuxt.js']:
        return 'Nuxt.js'
        
    return skill_name.title()


def extract_skills(text):
    """Extracts matching skills from the skill dictionary."""
    text_lower = text.lower()
    found_skills = []
    
    # Process text using spacy to filter tokens and do smart matching
    # 1. Regex word boundary matching for phrases and single words
    for skill in ALL_FLAT_SKILLS:
        # Avoid matching partial words (e.g., 'c' in 'concept' or 'go' in 'good')
        if len(skill) <= 2:
            pattern = r'\b' + re.escape(skill) + r'\b'
        else:
            pattern = r'\b' + re.escape(skill) + r'(?:s|es)?\b'
            
        if re.search(pattern, text_lower):
            found_skills.append(skill)
            
    # Clean duplicates and format nicely
    formatted_skills = []
    for skill in found_skills:
        formatted_skills.append(format_skill_name(skill))
                
    return list(set(formatted_skills))


def extract_sections(text):
    """Splits resume text into sections: Education, Experience, Certifications."""
    lines = text.split('\n')
    sections = {
        'Education': [],
        'Experience': [],
        'Certifications': []
    }
    
    education_keywords = ['education', 'academic background', 'qualification', 'academic credentials', 'academic profile']
    experience_keywords = ['experience', 'employment history', 'work history', 'professional background', 'career history', 'work experience', 'internship']
    cert_keywords = ['certification', 'certifications', 'licenses', 'credentials', 'courses', 'coursework']
    
    current_section = None
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check for headings
        line_lower = line_clean.lower()
        
        # Simple heading detection: line is short and matches keywords
        is_heading = False
        if len(line_clean) < 40:
            if any(kw == line_lower or line_lower.startswith(kw) for kw in experience_keywords):
                current_section = 'Experience'
                is_heading = True
            elif any(kw == line_lower or line_lower.startswith(kw) for kw in education_keywords):
                current_section = 'Education'
                is_heading = True
            elif any(kw == line_lower or line_lower.startswith(kw) for kw in cert_keywords):
                current_section = 'Certifications'
                is_heading = True
                
        if is_heading:
            continue
            
        # Add to section list
        if current_section:
            sections[current_section].append(line_clean)
            
    # Format paragraphs out of lists
    formatted_sections = {}
    for section_name, section_lines in sections.items():
        # Clean lines and group
        text_lines = []
        temp_para = []
        for l in section_lines:
            temp_para.append(l)
            if len(l) < 30 or l.endswith('.'):
                text_lines.append(" ".join(temp_para))
                temp_para = []
        if temp_para:
            text_lines.append(" ".join(temp_para))
            
        formatted_sections[section_name] = [p for p in text_lines if len(p.strip()) > 5]
        
    return formatted_sections


def parse_resume(file_path):
    """Full parsing orchestration for resumes."""
    text = extract_text(file_path)
    if not text:
        return None
        
    name = extract_name(text)
    email, phone = extract_contact_info(text)
    skills = extract_skills(text)
    sections = extract_sections(text)
    
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'education': sections.get('Education', []),
        'experience': sections.get('Experience', []),
        'certifications': sections.get('Certifications', []),
        'raw_text': text
    }


def parse_job_description(jd_text):
    """Parses a job description to extract required skills."""
    skills = extract_skills(jd_text)
    return {
        'skills': skills,
        'raw_text': jd_text
    }
